"""
Phase 2A – Template Filler
Fills the Word template per student and merges all into ONE .docx file,
with each student on a separate page.
"""

import os
import re
import copy
import zipfile
import shutil
import tempfile
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
import lxml.etree as etree

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import (
    OUTPUT_DIR, SUBJECTS, CDF_MAX_PER_SUBJECT, CDF_MAX_TOTAL,
    JEE_MAX_PER_SUBJECT, JEE_MAX_TOTAL, MAX_TESTS,
    ACADEMIC_YEAR,
)
from generate.chart_generator import compute_subject_averages


# ── Cell helpers ───────────────────────────────────────────────────────────────

def _first_real_cell(row_cells):
    """De-duplicate merged cells by underlying _tc element."""
    seen, unique = set(), []
    for c in row_cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            unique.append(c)
    return unique


def _set_cell_text(cell, text: str, bold: bool = True,
                   color: str = "000000", size: int = 20,
                   align=WD_ALIGN_PARAGRAPH.CENTER):
    """Clear a cell and write text with given formatting."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.size = Pt(size / 2)


def _add_text_to_cell(cell, label: str, value: str,
                      color: str = "001F60", size: int = 22):
    """Replace cell content with bold label + bold value."""
    para = cell.paragraphs[0]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para.clear()
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    for txt in (label, " " + value):
        run = para.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(r, g, b)


def _parse_test_number(exam_name: str) -> Optional[int]:
    """Extract test number from normalised exam name like 'JEE - 12' or 'CDF - 3'."""
    m = re.search(r"-\s*(\d+)\s*$", exam_name.strip())
    if m:
        return int(m.group(1))
    # Fallback: last number in string
    nums = re.findall(r"\d+", exam_name)
    return int(nums[-1]) if nums else None


# ── Chart PNG generation ───────────────────────────────────────────────────────

def _generate_chart_png(cdf_pcts: dict, jee_pcts: dict, output_path: str) -> None:
    """Generate a matplotlib percentage bar chart and save as PNG."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from config import CHART_COLORS

    subjects = SUBJECTS
    x = np.arange(len(subjects))
    width = 0.30

    fig, ax = plt.subplots(figsize=(4.3, 2.5), dpi=150)
    fig.patch.set_facecolor("white")

    cdf_vals = [cdf_pcts.get(s, 0) for s in subjects]
    jee_vals = [jee_pcts.get(s, 0) for s in subjects]

    bars1 = ax.bar(x - width / 2, cdf_vals, width,
                   label="CDF Avg %", color=CHART_COLORS["series1"],
                   edgecolor="white", linewidth=0.5, zorder=3)
    bars2 = ax.bar(x + width / 2, jee_vals, width,
                   label="JEE Avg %", color=CHART_COLORS["series2"],
                   edgecolor="white", linewidth=0.5, zorder=3)

    for bar in list(bars1) + list(bars2):
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, h + 0.8,
                    f"{h:.1f}%", ha="center", va="bottom",
                    fontsize=6.5, color="#333333")

    ax.set_xticks(x)
    ax.set_xticklabels(subjects, fontsize=8)
    ax.set_ylabel("Average %", fontsize=8)
    ax.set_ylim(0, 115)
    ax.set_yticks(range(0, 101, 20))
    ax.set_yticklabels([f"{v}%" for v in range(0, 101, 20)], fontsize=7)
    ax.set_title("")
    ax.legend(fontsize=7, loc="upper right")
    ax.yaxis.grid(True, linestyle="--", alpha=0.5, zorder=0)
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    plt.tight_layout(pad=0.4)
    plt.savefig(output_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ── Chart injection (PNG replaces chart1.xml) ──────────────────────────────────

# Exact dimensions of the chart drawing in the template (EMU units)
_CHART_CX = 3055108
_CHART_CY = 1782234

def _inline_image_xml(r_id: str) -> str:
    """Build the <w:drawing> XML for an inline PNG image with the given relationship ID."""
    return (
        f'<w:drawing'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{_CHART_CX}" cy="{_CHART_CY}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="ChartImage"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic>'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="1" name="ChartImage"/>'
        f'<pic:cNvPicPr/>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{r_id}"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{_CHART_CX}" cy="{_CHART_CY}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
    )


def _inject_chart_as_image(docx_path: str, cdf_pcts: dict, jee_pcts: dict) -> None:
    """
    Replace the embedded chart in a .docx with a matplotlib PNG image.
    Steps:
      1. Generate the PNG to a temp file.
      2. Open the .docx ZIP, add the PNG as word/media/chart_img.png.
      3. Add image relationship, remove chart relationship.
      4. Replace the <w:drawing>…<c:chart>…</w:drawing> block with inline image XML.
      5. Re-save the ZIP.
    """
    # 1. Generate PNG
    tmp_png = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    png_path = tmp_png.name
    tmp_png.close()
    try:
        _generate_chart_png(cdf_pcts, jee_pcts, png_path)
        with open(png_path, "rb") as f:
            png_bytes = f.read()
    finally:
        try:
            os.unlink(png_path)
        except OSError:
            pass

    img_rid      = "rIdChartImg"
    img_zip_path = "word/media/chart_img.png"
    img_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

    tmp_docx = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        doc_xml   = zin.read("word/document.xml").decode("utf-8")
        rels_xml  = zin.read("word/_rels/document.xml.rels").decode("utf-8")
        ct_xml    = zin.read("[Content_Types].xml").decode("utf-8")

        # 2. Replace chart drawing with inline image drawing
        new_drawing = _inline_image_xml(img_rid)
        doc_xml = re.sub(
            r"<w:drawing>(?:(?!<w:drawing>).)*?<c:chart(?:(?!<w:drawing>).)*?</w:drawing>",
            new_drawing,
            doc_xml,
            flags=re.DOTALL,
        )

        # 3. Remove old chart + embedding relationships, add new image relationship
        rels_xml = re.sub(
            r'<Relationship[^>]+Target="(?:charts|embeddings)/[^"]*"[^>]*/>', "", rels_xml
        )
        new_img_rel = (
            f'<Relationship Id="{img_rid}" Type="{img_rel_type}"'
            f' Target="media/chart_img.png"/>'
        )
        rels_xml = rels_xml.replace("</Relationships>", new_img_rel + "</Relationships>")

        # 4. Remove chart + embedding Override entries; add PNG content type
        ct_xml = re.sub(
            r'<Override\s+PartName="/word/(?:charts|embeddings)/[^"]*"[^>]*/>', "", ct_xml
        )
        if 'Extension="png"' not in ct_xml:
            ct_xml = ct_xml.replace(
                "</Types>",
                '<Default Extension="png" ContentType="image/png"/></Types>',
            )

        with zipfile.ZipFile(tmp_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                fname = item.filename
                if fname == "word/document.xml":
                    zout.writestr(item, doc_xml.encode("utf-8"))
                elif fname == "word/_rels/document.xml.rels":
                    zout.writestr(item, rels_xml.encode("utf-8"))
                elif fname == "[Content_Types].xml":
                    zout.writestr(item, ct_xml.encode("utf-8"))
                elif fname.startswith("word/charts/") or fname.startswith("word/embeddings/"):
                    pass  # drop chart XML and embedded Excel
                else:
                    zout.writestr(item, zin.read(fname))
            # 5. Add PNG to the archive
            zout.writestr(img_zip_path, png_bytes)

    os.replace(tmp_docx, docx_path)


# ── Merge helper ───────────────────────────────────────────────────────────────

def _extract_body_content(doc_xml: str) -> str:
    """
    Return the inner XML of <w:body> excluding the trailing <w:sectPr> element.
    This is the content that gets appended for each subsequent student.
    """
    m = re.search(r"<w:body>(.*)</w:body>", doc_xml, re.DOTALL)
    if not m:
        return ""
    body = m.group(1)
    # Remove the section-properties block that ends the body
    body = re.sub(r"<w:sectPr\b.*?</w:sectPr>\s*$", "", body, flags=re.DOTALL)
    return body.strip()


_PAGE_BREAK_XML = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    "<w:r><w:br w:type=\"page\"/></w:r>"
    "</w:p>"
)


def _build_master_docx(temp_paths: list[str], output_path: str) -> None:
    """
    Merge a list of per-student .docx files (each with a PNG chart image)
    into one master .docx with a page break between students.

    Works at the ZIP level so image relationships are handled correctly.
    """
    # Read first student's full archive as the base
    with zipfile.ZipFile(temp_paths[0], "r") as z0:
        base = {item.filename: z0.read(item.filename) for item in z0.infolist()}

    master_doc_xml  = base["word/document.xml"].decode("utf-8")
    master_rels_xml = base["word/_rels/document.xml.rels"].decode("utf-8")

    img_counter = 100  # start high to avoid clashing with template's existing IDs

    for idx, path in enumerate(temp_paths[1:], start=1):
        with zipfile.ZipFile(path, "r") as z:
            sub_doc_xml  = z.read("word/document.xml").decode("utf-8")
            sub_rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")

            # Find image relationships in this student's doc
            img_rels = {}
            for m in re.finditer(
                r'<Relationship\s[^>]*Id="([^"]+)"[^>]*Type="([^"]*image[^"]*)"[^>]*Target="([^"]+)"',
                sub_rels_xml,
            ):
                img_rels[m.group(1)] = (m.group(2), m.group(3))

            # Copy image files into master, assign new relationship IDs.
            # The chart image (chart_img.png) is ALWAYS unique per student.
            # Shared images (e.g. the school logo) that already exist in base
            # are reused without copying.
            rId_remap: dict[str, str] = {}
            for old_rid, (rel_type, target) in img_rels.items():
                zip_path   = f"word/{target}"
                is_chart   = "chart_img" in target   # unique per student

                if not is_chart and zip_path in base:
                    # Shared image (e.g. logo) – find its existing rId in master rels
                    existing = re.search(
                        rf'<Relationship\s[^>]*Id="([^"]+)"[^>]*Target="{re.escape(target)}"',
                        master_rels_xml,
                    )
                    if existing:
                        rId_remap[old_rid] = existing.group(1)
                    continue

                img_bytes = z.read(zip_path)
                ext       = target.rsplit(".", 1)[-1]
                new_name  = f"media/chart_img_{img_counter}.{ext}"
                new_rid   = f"rIdChartImg{img_counter}"
                img_counter += 1

                base[f"word/{new_name}"] = img_bytes
                master_rels_xml = master_rels_xml.replace(
                    "</Relationships>",
                    f'<Relationship Id="{new_rid}" Type="{rel_type}"'
                    f' Target="{new_name}"/></Relationships>',
                )
                rId_remap[old_rid] = new_rid

            # Remap relationship IDs in this student's body XML
            sub_doc_xml_remapped = sub_doc_xml
            for old, new in rId_remap.items():
                sub_doc_xml_remapped = sub_doc_xml_remapped.replace(
                    f'r:embed="{old}"', f'r:embed="{new}"'
                )

            # Extract body content and append (with page break) into master
            body_content = _extract_body_content(sub_doc_xml_remapped)
            master_doc_xml = master_doc_xml.replace(
                "</w:body>",
                _PAGE_BREAK_XML + body_content + "</w:body>",
            )

    # Update master relationships XML in base dict
    base["word/document.xml"]          = master_doc_xml.encode("utf-8")
    base["word/_rels/document.xml.rels"] = master_rels_xml.encode("utf-8")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for fname, data in base.items():
            zout.writestr(fname, data)


# ── Main fill function ─────────────────────────────────────────────────────────

def fill_template(
    student_username: str,
    student_name: str,
    batch: str,
    exams_data: list[dict],
    template_path: str,
    output_docx_path: str,
) -> None:
    """
    Fill the Word template for one student and save to output_docx_path.
    Chart is embedded as a matplotlib PNG (not chart XML) so the file
    can be safely merged into a multi-student document.
    """
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)

    doc   = Document(template_path)
    table = doc.tables[0]

    # ── Header rows ───────────────────────────────────────────────────────────
    row2 = _first_real_cell(table.rows[2].cells)
    if row2:
        _add_text_to_cell(row2[0], "Name of the Student:", student_name)
    if len(row2) > 1:
        _add_text_to_cell(row2[1], "Grade:", batch)

    row3 = _first_real_cell(table.rows[3].cells)
    if row3:
        _add_text_to_cell(row3[0], "User Id:", student_username)
    if len(row3) > 1:
        _add_text_to_cell(row3[1], "A.Y :", ACADEMIC_YEAR)

    # ── Marks table rows ──────────────────────────────────────────────────────
    # Separate CDF and JEE records into independent dicts keyed by test number.
    # Each exam is one type only — CDF scores fill cols 1-4, JEE scores fill
    # cols 5-8. No scaling between types.
    cdf_records: dict[int, dict] = {}
    jee_records: dict[int, dict] = {}

    for rec in exams_data:
        ename = rec.get("Exam", "").strip()
        tnum  = _parse_test_number(ename)
        if not tnum or not (1 <= tnum <= MAX_TESTS):
            continue
        if ename.upper().startswith("CDF"):
            cdf_records[tnum] = rec
        else:
            jee_records[tnum] = rec

    def to_int(v):
        s = str(v).strip()
        if s.upper() == "AB":
            return "AB"
        if s == "-":
            return "-"
        try:
            return int(float(s))
        except (ValueError, TypeError):
            return ""

    def fill_score_cols(row_cells, math_v, phy_v, chem_v, col_start):
        """Write math/phy/chem and total into cols col_start … col_start+3."""
        vals = [math_v, phy_v, chem_v]
        if any(v == "AB" for v in vals):
            total = "AB"
        elif any(v == "-" for v in vals):
            total = "-"
        elif all(isinstance(v, int) for v in vals):
            total = math_v + phy_v + chem_v
        else:
            total = ""
        for col_offset, val in enumerate([math_v, phy_v, chem_v, total]):
            col_idx = col_start + col_offset
            if col_idx < len(row_cells) and val != "":
                _set_cell_text(
                    row_cells[col_idx],
                    str(val),
                    bold=True,
                    color="000000",
                    size=20,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )

    for test_num in range(1, MAX_TESTS + 1):
        row_idx = 5 + test_num          # row 6 = Test-1, row 7 = Test-2 …
        if row_idx >= len(table.rows):
            break

        row_cells = _first_real_cell(table.rows[row_idx].cells)

        # CDF columns (1-4): Math CDF | Phy CDF | Chem CDF | Total CDF
        if test_num in cdf_records:
            rec = cdf_records[test_num]
            fill_score_cols(
                row_cells,
                to_int(rec.get("Mathematics Score", "")),
                to_int(rec.get("Physics Score", "")),
                to_int(rec.get("Chemistry Score", "")),
                col_start=1,
            )

        # JEE columns (5-8): Math JEE | Phy JEE | Chem JEE | Total JEE
        if test_num in jee_records:
            rec = jee_records[test_num]
            fill_score_cols(
                row_cells,
                to_int(rec.get("Mathematics Score", "")),
                to_int(rec.get("Physics Score", "")),
                to_int(rec.get("Chemistry Score", "")),
                col_start=5,
            )

    # ── Remark (Row 20) ──────────────────────────────────────────────────────
    cdf_pcts, jee_pcts = compute_subject_averages(exams_data)

    # Detect all-AB: student has exam records but every score was absent
    has_any_numeric = any(
        str(rec.get(col, "")).strip().upper() != "AB"
        and str(rec.get(col, "")).strip() != ""
        for rec in exams_data
        for col in ["Mathematics Score", "Physics Score", "Chemistry Score"]
    )
    all_ab = not has_any_numeric and len(exams_data) > 0

    if all_ab:
        remark = "Insufficient data due to absences. Regular attendance is strongly advised."
    else:
        # CDF overall % = mean of non-zero subject percentages (only if exams exist)
        cdf_vals = [v for v in cdf_pcts.values() if v > 0]
        jee_vals = [v for v in jee_pcts.values() if v > 0]

        overalls = []
        if cdf_vals:
            overalls.append(sum(cdf_vals) / len(cdf_vals))
        if jee_vals:
            overalls.append(sum(jee_vals) / len(jee_vals))

        combined_pct = sum(overalls) / len(overalls) if overalls else 0.0

        if combined_pct >= 90:
            remark = "Outstanding performance! Keep up the excellent work."
        elif combined_pct >= 75:
            remark = "Great work! A little more focus on weak areas will make you unstoppable."
        elif combined_pct >= 60:
            remark = "Good effort. Needs to strengthen core concepts for better results."
        elif combined_pct >= 45:
            remark = "Needs improvement. Daily practice and focused study required."
        else:
            remark = "Needs significant improvement. Immediate attention and extra support advised."

    remarks_row = _first_real_cell(table.rows[17].cells)
    if remarks_row:
        _add_text_to_cell(remarks_row[0], "Remarks:", remark)

    # ── Save interim .docx ───────────────────────────────────────────────────
    doc.save(output_docx_path)

    # ── Replace chart with matplotlib PNG ───────────────────────────────────
    _inject_chart_as_image(output_docx_path, cdf_pcts, jee_pcts)


# ── Batch generation ──────────────────────────────────────────────────────────

def generate_report_cards(
    csv_path: str,
    template_path: str,
    batch_filter: Optional[str] = None,
) -> str | None:
    """
    Generate one merged .docx file containing all students (one page each).
    Returns the path to the merged file, or None if nothing was generated.
    """
    import csv as _csv

    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"CSV not found: {csv_path}. Run 'ingest' first.")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    with open(csv_path, newline="", encoding="utf-8") as f:
        all_rows = list(_csv.DictReader(f))

    if not all_rows:
        print("CSV is empty. Nothing to generate.")
        return None

    # Group rows by student.
    # Pass 1: find which usernames belong to the selected batch.
    # Pass 2: collect ALL records for those usernames (CDF and JEE exams
    #         may carry different batch labels, so we must not filter by
    #         batch when collecting exam records — only when deciding which
    #         students to include).
    if batch_filter:
        eligible_uids: set[str] = {
            row["Username"].strip()
            for row in all_rows
            if row.get("Batch", "").strip() == batch_filter
            and row.get("Username", "").strip()
        }
    else:
        eligible_uids = {
            row["Username"].strip()
            for row in all_rows
            if row.get("Username", "").strip()
        }

    students: dict[str, list[dict]] = {}
    for row in all_rows:
        uid = row.get("Username", "").strip()
        if uid and uid in eligible_uids:
            students.setdefault(uid, []).append(row)

    if not students:
        msg = f" for batch '{batch_filter}'" if batch_filter else ""
        print(f"No students found{msg}.")
        return None

    total = len(students)
    print(f"\nGenerating report cards for {total} student(s)…")

    # Work in a temp directory; final merge goes to output/docx/
    tmp_dir  = tempfile.mkdtemp(prefix="rc_tmp_")
    docx_dir = os.path.join(OUTPUT_DIR, "docx")
    os.makedirs(docx_dir, exist_ok=True)

    temp_paths: list[str] = []

    try:
        for i, (username, exams) in enumerate(students.items(), 1):
            first_name = exams[0].get("First Name", username)
            # Use the batch from the selected filter; if no filter, use the
            # first CDF record's batch (primary batch), falling back to first record.
            if batch_filter:
                batch = batch_filter
            else:
                primary = next(
                    (r for r in exams if r.get("Exam", "").upper().startswith("CDF")),
                    exams[0],
                )
                batch = primary.get("Batch", "")
            safe_name  = re.sub(r'[\\/:*?"<>|]', "_", username)
            tmp_path   = os.path.join(tmp_dir, f"{safe_name}.docx")

            print(f"  [{i}/{total}] {first_name} ({username}) — {len(exams)} exam(s)")
            try:
                fill_template(
                    student_username=username,
                    student_name=first_name,
                    batch=batch,
                    exams_data=exams,
                    template_path=template_path,
                    output_docx_path=tmp_path,
                )
                temp_paths.append(tmp_path)
            except Exception as exc:
                print(f"  [ERROR] {username}: {exc}")

        if not temp_paths:
            print("No report cards generated.")
            return None

        # Determine output filename
        tag      = re.sub(r'[\\/:*?"<>|]', "_", batch_filter) if batch_filter else "ALL"
        out_path = os.path.join(docx_dir, f"{tag}_report_cards.docx")

        print(f"\nMerging {len(temp_paths)} student(s) into one file…")
        _build_master_docx(temp_paths, out_path)
        print(f"Done. Saved: {out_path}")
        return out_path

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
